#!/usr/bin/env python3
import os, sys, json, base64, requests, argparse, re, io, shutil, tempfile
from typing import Dict, Any, List, Tuple, Optional
from pathlib import Path

from dotenv import load_dotenv
load_dotenv()

# ==== OCR config ==============================================================
MISTRAL_OCR_ENDPOINT = os.environ.get("MISTRAL_OCR_ENDPOINT")
MISTRAL_API_KEY      = os.environ.get("API_KEY")
MISTRAL_MODEL        = os.environ.get("MISTRAL_MODEL", "mistral-document-ai-2505")

# ==== Optional Pandoc (for LaTeX math & tables) ===============================
HAVE_PYPANDOC = True
try:
    import pypandoc  # type: ignore
except Exception:
    HAVE_PYPANDOC = False

def ensure_pandoc_available() -> bool:
    if shutil.which("pandoc"):
        return True
    if HAVE_PYPANDOC:
        try:
            pypandoc.download_pandoc()
            return True
        except Exception:
            return False
    return False

# ==== Imaging / PDF / DOCX fallback ==========================================
from PIL import Image
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Inches, Pt

# ==== Basics =================================================================
def die(msg: str, code: int = 1):
    print(f"[ERR] {msg}", file=sys.stderr); sys.exit(code)

def post_ocr(payload: dict):
    if not MISTRAL_OCR_ENDPOINT or not MISTRAL_API_KEY:
        die("Set MISTRAL_OCR_ENDPOINT and MISTRAL_API_KEY in your environment (.env).")
    headers = {"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"}
    print(f"[POST] {MISTRAL_OCR_ENDPOINT}")
    r = requests.post(MISTRAL_OCR_ENDPOINT, headers=headers, json=payload, timeout=600)
    print(f"[HTTP] {r.status_code} ctype={r.headers.get('content-type')}")
    if r.status_code >= 400:
        print("[BODY]", r.text[:4000])
        r.raise_for_status()
    return r.json()

def bytes_to_data_url(mime: str, data: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(data).decode('utf-8')}"

def _strip(s: str) -> str:
    return (s or "").strip()

def unwrap_container(resp: Dict[str, Any]) -> Dict[str, Any]:
    node = resp
    for k in ("output","response","result","data","document"):
        if isinstance(node, dict) and isinstance(node.get(k), dict):
            node = node[k]
    return node

# ==== Extract text from OCR pages ============================================
def extract_from_page(p: Dict[str, Any]) -> str:
    md = p.get("markdown")
    if isinstance(md, str) and md.strip():
        return md.strip()

    candidates: List[str] = []

    if isinstance(p.get("lines"), list):
        parts = []
        for ln in p["lines"]:
            if isinstance(ln, dict):
                t = _strip(ln.get("content") or ln.get("text"))
                if t: parts.append(t)
        if parts: candidates.append("\n".join(parts))

    if isinstance(p.get("paragraphs"), list):
        parts = []
        for para in p["paragraphs"]:
            if isinstance(para, dict):
                t = _strip(para.get("content") or para.get("text"))
                if t: parts.append(t)
        if parts: candidates.append("\n".join(parts))

    for key in ("blocks","items","elements","regions"):
        arr = p.get(key)
        if isinstance(arr, list) and arr:
            parts = []
            for it in arr:
                if isinstance(it, dict):
                    t = _strip(it.get("text") or it.get("content") or it.get("value"))
                    if t: parts.append(t)
            if parts: candidates.append("\n".join(parts))

    t = _strip(p.get("content") or p.get("text") or p.get("full_text") or p.get("raw_text"))
    if t: candidates.append(t)

    candidates = [c for c in candidates if _strip(c)]
    return max(candidates, key=len) if candidates else ""

# ==== Image I/O from OCR ======================================================
_DATA_URL_RE = re.compile(r'^data:(?P<mime>[^;]+);base64,(?P<b64>.+)$', re.IGNORECASE)

def save_image_bytes_clean(raw: bytes, out_path: Path) -> Path:
    """Normalize to PNG/JPEG; RGB; return final path with correct ext."""
    with Image.open(io.BytesIO(raw)) as im:
        fmt = (im.format or "").upper()
        if fmt == "PNG" and im.mode == "RGB":
            out = out_path.with_suffix(".png")
            im.save(out)
            return out
        im = im.convert("RGB")
        out = out_path.with_suffix(".jpg")
        im.save(out, quality=92, optimize=True)
        return out

def save_base64_image_unknown(input_str: str, out_dir: Path, stem: str) -> Path:
    """
    Accepts either a bare base64 string or a full data: URL and saves a normalized image.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    s = input_str.strip()
    m = _DATA_URL_RE.match(s)
    if m:
        raw = base64.b64decode(m.group("b64"), validate=False)  # <-- correct 'b64'
    else:
        # bare b64 OR string that still contains 'base64,'
        try:
            raw = base64.b64decode(s, validate=False)
        except Exception:
            raw = base64.b64decode(s.split("base64,", 1)[-1], validate=False)
    return save_image_bytes_clean(raw, out_dir / stem)

def fetch_and_save_http_image(url: str, out_dir: Path, stem: str) -> Optional[Path]:
    """
    If OCR returns http(s) image URLs instead of base64/data URLs, fetch & save.
    """
    try:
        r = requests.get(url, timeout=60)
        r.raise_for_status()
        raw = r.content
        try:
            with Image.open(io.BytesIO(raw)) as im:
                fmt = (im.format or "").upper()
                if fmt == "PNG" and im.mode == "RGB":
                    out = out_dir / f"{stem}.png"
                    im.save(out)
                    return out
                im = im.convert("RGB")
                out = out_dir / f"{stem}.jpg"
                im.save(out, quality=92, optimize=True)
                return out
        except Exception:
            out = out_dir / f"{stem}.bin"
            out.write_bytes(raw)
            return out
    except Exception as e:
        print(f"[WARN] fetch failed {url}: {e}")
        return None

# ==== PDF render + crops ======================================================
def render_pdf_page_to_image(pdf_path: Path, page_num: int, dpi: int = 300) -> Tuple[Image.Image, Tuple[float,float]]:
    doc = fitz.open(pdf_path.as_posix())
    try:
        page = doc[page_num-1]
        zoom = dpi / 72.0
        pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img  = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        return img, (page.rect.width, page.rect.height)
    finally:
        doc.close()

def clamp(v, lo, hi): return max(lo, min(hi, v))

def bbox_to_pixels(b: Tuple[float,float,float,float],
                   img_w: int, img_h: int,
                   pts_wh: Optional[Tuple[float,float]] = None,
                   coord_type: str = "norm",
                   origin: str = "top-left") -> Tuple[int,int,int,int]:
    x0,y0,x1,y1 = b
    if coord_type == "norm":
        if origin == "bottom-left":
            y0, y1 = 1 - y0, 1 - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0*img_w)), int(round(y0*img_h)), int(round(x1*img_w)), int(round(y1*img_h))
    elif coord_type == "pixel":
        if origin == "bottom-left":
            y0, y1 = img_h - y0, img_h - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0)), int(round(y0)), int(round(x1)), int(round(y1))
    elif coord_type == "pdf_points":
        if not pts_wh: raise ValueError("pdf_points requires page size in points")
        pts_w, pts_h = pts_wh; sx, sy = img_w/pts_w, img_h/pts_h
        if origin == "bottom-left":
            y0, y1 = pts_h - y0, pts_h - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0*sx)), int(round(y0*sy)), int(round(x1*sx)), int(round(y1*sy))
    else:
        raise ValueError("coord_type must be norm|pixel|pdf_points")
    x0p,x1p = sorted([X0,X1]); y0p,y1p = sorted([Y0,Y1])
    x0p = clamp(x0p,0,img_w-1); x1p = clamp(x1p,1,img_w)
    y0p = clamp(y0p,0,img_h-1); y1p = clamp(y1p,1,img_h)
    return x0p,y0p,x1p,y1p

def crop_and_save(img: Image.Image, bbox_px: Tuple[int,int,int,int],
                  padding: int, out_dir: Path, stem: str) -> Path:
    x0,y0,x1,y1 = bbox_px
    if padding:
        x0 = max(0, x0-padding); y0 = max(0, y0-padding)
        x1 = min(img.width, x1+padding); y1 = min(img.height, y1+padding)
    crop = img.crop((x0,y0,x1,y1)).convert("RGB")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{stem}.jpg"
    crop.save(out_path, quality=92, optimize=True)
    return out_path

# Always-on: if no regions.json provided, generate a full-page crop per page
def autocrops_full_pages(pdf_path: Path, dpi: int, assets_dir: Path) -> Dict[int, List[Path]]:
    out: Dict[int, List[Path]] = {}
    doc = fitz.open(pdf_path.as_posix())
    try:
        for i in range(len(doc)):
            page_num = i + 1
            img, _ = render_pdf_page_to_image(pdf_path, page_num, dpi=dpi)
            saved = crop_and_save(img, (0, 0, img.width, img.height), 0, assets_dir, f"p{page_num}_full")
            out.setdefault(page_num, []).append(saved)
    finally:
        doc.close()
    print(f"[INFO] auto-crops (full pages) generated: {sum(len(v) for v in out.values())}")
    return out

# ==== Markdown helpers ========================================================
def clean_markdown(md: str) -> str:
    md = md.replace("\r\n", "\n").replace("\r", "\n")
    return md

_MATH_OR_TABLE_RE = re.compile(
    r"(\$\$.*?\$\$|\$[^$\n]+\$|\\\(|\\\)|\\\[|\\\]|\\begin\{(equation|align|eqnarray|gather|aligned)\}|(^\s*\|.*\|\s*$\n^\s*\|?\s*[-:]+\s*(\|[-:]+\s*)+$))",
    re.MULTILINE | re.DOTALL
)

def detect_math_or_tables(pages_text: List[str]) -> bool:
    return bool(_MATH_OR_TABLE_RE.search("\n\n".join(pages_text)))

_IMG_MD_RE   = re.compile(r'!\[[^\]]*\]\([^)]+\)', re.IGNORECASE)
_IMG_HTML_RE = re.compile(r'<img\b[^>]*>', re.IGNORECASE)

def strip_inline_images(md: str) -> str:
    return _IMG_HTML_RE.sub('', _IMG_MD_RE.sub('', md))

def md_image(path: Path, width_in: float) -> str:
    p = Path(path).resolve().as_posix()
    return f'![]({p}){{width={width_in}in}}'

def build_markdown(pages_text: List[str],
                   images_by_page: Dict[int, List[Path]],
                   crops_by_page: Dict[int, List[Path]],
                   insert_page_breaks: bool,
                   image_max_width_in: float) -> str:
    parts: List[str] = []
    for i, txt in enumerate(pages_text, start=1):
        text_clean = strip_inline_images(_strip(txt))
        parts.append(f"\n\n## Page {i}\n\n{text_clean}\n")
        for im in images_by_page.get(i, []):
            parts.append("\n" + md_image(im, image_max_width_in) + "\n")
        for im in crops_by_page.get(i, []):
            parts.append("\n" + md_image(im, image_max_width_in) + "\n")
        if insert_page_breaks and i < len(pages_text):
            parts.append("\n\\newpage\n")
    return clean_markdown("".join(parts)).strip() + "\n"

def build_docx_with_pandoc_to_path(md_text: str, out_path: Path, resource_dirs: Optional[List[Path]] = None) -> None:
    if not HAVE_PYPANDOC:
        raise RuntimeError("pypandoc not installed")
    if not ensure_pandoc_available():
        raise RuntimeError("Pandoc not available and auto-download failed")
    with tempfile.TemporaryDirectory() as td:
        md_file = Path(td) / "in.md"
        md_file.write_text(md_text, encoding="utf-8")
        extra_args = ["--standalone"]
        if resource_dirs:
            search_path = os.pathsep.join(str(Path(p).resolve()) for p in resource_dirs)
            extra_args.append(f"--resource-path={search_path}")
        pypandoc.convert_file(
            str(md_file),
            to="docx",
            format="gfm+tex_math_dollars+pipe_tables",
            outputfile=str(out_path),
            extra_args=extra_args,
        )

# ==== Basic DOCX fallback =====================================================
def add_text_block(doc: Document, text: str):
    for chunk in text.split("\n\n"):
        chunk = chunk.strip()
        if chunk:
            doc.add_paragraph(chunk)

def add_picture_fit_width(doc: Document, image_path: Path, max_width_in: float):
    doc.add_picture(str(image_path), width=Inches(max_width_in))

def build_docx_with_python_docx_to_path(pages_text: List[str],
                                        images_by_page: Dict[int, List[Path]],
                                        crops_by_page: Dict[int, List[Path]],
                                        out_path: Path,
                                        insert_page_breaks: bool,
                                        image_max_width_in: float) -> None:
    doc = Document()
    style = doc.styles['Normal'].font
    style.name = "Calibri"; style.size = Pt(11)
    for i, txt in enumerate(pages_text, start=1):
        doc.add_heading(f"Page {i}", level=2)
        if _strip(txt):
            add_text_block(doc, txt)
        for im in images_by_page.get(i, []):
            add_picture_fit_width(doc, im, image_max_width_in)
        for im in crops_by_page.get(i, []):
            add_picture_fit_width(doc, im, image_max_width_in)
        if insert_page_breaks and i < len(pages_text):
            doc.add_page_break()
    doc.save(out_path)

# ==== Regions JSON -> saved crops ============================================
def crops_from_regions(pdf_path: Path, regions_json_path: Path, dpi: int, assets_dir: Path) -> Dict[int, List[Path]]:
    with open(regions_json_path, "r", encoding="utf-8") as jf:
        cfg = json.load(jf)
    pages_cfg = cfg.get("pages", {})
    out: Dict[int, List[Path]] = {}
    for p_str, regs in pages_cfg.items():
        try:
            pnum = int(p_str)
        except:
            continue
        if not isinstance(regs, list):
            continue
        page_img, pts_wh = render_pdf_page_to_image(pdf_path, pnum, dpi=dpi)
        for k, r in enumerate(regs, start=1):
            coords     = r["coords"]
            coord_type = r.get("coord_type","norm")
            origin     = r.get("origin","top-left")
            padding    = int(r.get("padding", 8))
            label      = r.get("label", f"crop{k}")
            bbox_px = bbox_to_pixels(tuple(coords), page_img.width, page_img.height,
                                     pts_wh if coord_type=="pdf_points" else None,
                                     coord_type=coord_type, origin=origin)
            saved = crop_and_save(page_img, bbox_px, padding, assets_dir, f"p{pnum}_{k}_{label}")
            out.setdefault(pnum, []).append(saved)
    return out

# ==== Output path resolver (robust) ===========================================
def resolve_output_paths(pdf_path: Path, out_arg: Optional[str]) -> Tuple[Path, Path]:
    if out_arg:
        out = Path(out_arg)
        if out.suffix.lower() == ".docx":
            docx_path = out
            assets_dir = out.parent / (out.stem + "_assets")
        elif out.exists() and out.is_dir():
            docx_path = out / (pdf_path.stem + "_ocr.docx")
            assets_dir = out / (pdf_path.stem + "_assets")
        else:
            docx_path = Path(str(out) + "_ocr.docx")
            assets_dir = Path(str(out) + "_assets")
    else:
        docx_path = pdf_path.with_name(pdf_path.stem + "_ocr.docx")
        assets_dir = pdf_path.with_name(pdf_path.stem + "_assets")
    return docx_path, assets_dir

# ==== Helper to collect images (recursive + tolerant) =========================
def _iter_possible_images(p: Dict[str, Any]) -> List[Dict[str, Any]]:
    imgs: List[Dict[str, Any]] = []
    for key in ("images", "figures", "media", "inline_images"):
        val = p.get(key)
        if isinstance(val, list):
            imgs.extend([v for v in val if isinstance(v, dict)])

    def rec(node):
        if isinstance(node, dict):
            maybe_url = node.get("url") or node.get("src") or node.get("image")
            maybe_b64 = node.get("base64") or node.get("image_base64") or node.get("data") or node.get("imageData") or node.get("content") or node.get("b64")
            if isinstance(maybe_b64, str) and (len(maybe_b64) > 100 or maybe_b64.lower().startswith("data:")):
                imgs.append(node)
            elif isinstance(maybe_url, str) and (maybe_url.lower().startswith("data:") or maybe_url.lower().startswith("http")):
                imgs.append(node)
            for v in node.values():
                rec(v)
        elif isinstance(node, list):
            for v in node:
                rec(v)
    rec(p)

    seen = set()
    deduped = []
    for im in imgs:
        key = json.dumps({k: im.get(k) for k in ("id","url","src","image")}, sort_keys=True, default=str)
        if key not in seen:
            seen.add(key)
            deduped.append(im)
    return deduped

# ==== Main ====================================================================
def main():
    if not MISTRAL_OCR_ENDPOINT or not MISTRAL_API_KEY:
        die("Set MISTRAL_OCR_ENDPOINT and MISTRAL_API_KEY in your environment (.env).")

    ap = argparse.ArgumentParser(description="OCR PDF → DOCX (Pandoc for math) with images & crops (regions or auto).")
    ap.add_argument("pdf", nargs="+", help="Path to the PDF (quote if it has spaces)")
    ap.add_argument("--title", default=None, help="DOCX title heading (markdown H1)")
    ap.add_argument("--out", default=None, help="Base output name/path (prefix, dir, or final .docx)")
    ap.add_argument("--regions-json", default=None, help="Crop & embed regions from JSON (always used when provided)")
    ap.add_argument("--dpi", type=int, default=300, help="Render DPI for cropping")
    ap.add_argument("--image-max-width", type=float, default=6.5, help="Image width in inches")
    ap.add_argument("--no-page-breaks", dest="no_page_breaks", action="store_true", help="No page breaks between pages")
    ap.add_argument("--prefer-pandoc", action="store_true", help="Prefer Pandoc if available")
    ap.add_argument("--force-pandoc", action="store_true", help="Force Pandoc; error if unavailable")
    args = ap.parse_args()

    path = " ".join(args.pdf)
    pdf_path = Path(path)
    if not pdf_path.exists():
        die(f"File not found: {pdf_path}")

    pdf_bytes = pdf_path.read_bytes()
    print(f"[INFO] file={pdf_path.name} bytes={len(pdf_bytes)}")

    payload = {
        "model": MISTRAL_MODEL,
        "document": {"type": "document_url", "document_url": bytes_to_data_url("application/pdf", pdf_bytes)},
        "include_image_base64": True,  # ALWAYS ask OCR to include images if it can
    }
    resp = post_ocr(payload)
    Path("ocr_response.json").write_text(json.dumps(resp, ensure_ascii=False, indent=2), encoding="utf-8")
    print("[INFO] wrote ocr_response.json")

    container = unwrap_container(resp)
    pages = container.get("pages")
    if not isinstance(pages, list) or not pages:
        top = ""
        for k in ("markdown","full_text","content","text","raw_text"):
            if isinstance(container.get(k), str) and container[k].strip():
                top = container[k]; break
        if not top.strip():
            die("No pages and no usable top-level text found in OCR response.")
        pages = [{"markdown": top}]

    docx_path, assets_dir = resolve_output_paths(pdf_path, args.out)
    print(f"[OUT] docx={docx_path}")
    print(f"[OUT] assets_dir={assets_dir}")

    if assets_dir.exists():
        shutil.rmtree(assets_dir, ignore_errors=True)
    assets_dir.mkdir(parents=True, exist_ok=True)

    pages_text: List[str] = []
    ocr_images_by_page: Dict[int, List[Path]] = {}

    print(f"[OK] pages={len(pages)}")
    for i, p in enumerate(pages, start=1):
        txt = extract_from_page(p if isinstance(p, dict) else {}) or ""
        if args.title and i == 1:
            txt = f"# {args.title}\n\n{txt}"
        pages_text.append(txt)

        # ALWAYS try to collect images the OCR provides
        if isinstance(p, dict):
            ims = _iter_possible_images(p)
            print(f"[INFO] page {i}: found {len(ims)} potential image nodes")
            for j, im in enumerate(ims, start=1):
                saved = None
                # base64-like fields
                for k in ("base64","image_base64","data","imageData","content","b64"):
                    v = im.get(k)
                    if isinstance(v, str) and (len(v) > 100 or v.lower().startswith("data:")):
                        try:
                            saved = save_base64_image_unknown(v, assets_dir, f"p{i}_{j}_ocrimg")
                            print(f"[OK] saved OCR image (b64): {saved.name}")
                            break
                        except Exception as e:
                            print(f"[WARN] b64 save failed: {e}")
                # data-URL / http(s) URL fields
                if not saved:
                    for k in ("data_url","url","image","src"):
                        v = im.get(k)
                        if isinstance(v, str) and v.strip():
                            if v.lower().startswith("data:"):
                                try:
                                    saved = save_base64_image_unknown(v, assets_dir, f"p{i}_{j}_ocrimg")
                                    print(f"[OK] saved OCR image (data-url): {saved.name}")
                                    break
                                except Exception as e:
                                    print(f"[WARN] data-url save failed: {e}")
                            elif v.lower().startswith(("http://","https://")):
                                saved = fetch_and_save_http_image(v, assets_dir, f"p{i}_{j}_ocrimg")
                                if saved:
                                    print(f"[OK] fetched OCR image (http): {saved.name}")
                                    break
                if saved:
                    ocr_images_by_page.setdefault(i, []).append(saved)

    total_imgs = sum(len(v) for v in ocr_images_by_page.values())
    print(f"[INFO] total OCR images saved: {total_imgs}")

    # === ALWAYS INCLUDE REGION CROPS ===
    # If user supplied regions.json, use it; else auto-crop full page per page.
    if args.regions_json:
        rj = Path(args.regions_json)
        if not rj.exists():
            die(f"regions JSON not found: {rj}")
        crops_by_page = crops_from_regions(pdf_path, rj, args.dpi, assets_dir)
        total_crops = sum(len(v) for v in crops_by_page.values())
        print(f"[INFO] saved {total_crops} crop(s) from regions (user JSON)")
    else:
        crops_by_page = autocrops_full_pages(pdf_path, args.dpi, assets_dir)

    wants_math = detect_math_or_tables(pages_text)
    prefer_pd  = args.prefer_pandoc or args.force_pandoc or wants_math

    try:
        if prefer_pd:
            if not ensure_pandoc_available():
                if args.force_pandoc:
                    die("Pandoc required (--force-pandoc) but not available.")
                else:
                    print("[WARN] Pandoc unavailable; falling back to basic DOCX.")
                    raise RuntimeError("no-pandoc")
            md_text = build_markdown(
                pages_text,
                images_by_page=ocr_images_by_page,
                crops_by_page=crops_by_page,
                insert_page_breaks=not args.no_page_breaks,
                image_max_width_in=args.image_max_width
            )
            build_docx_with_pandoc_to_path(md_text, docx_path, resource_dirs=[assets_dir])
            print(f"[OK] wrote docx (pandoc): {docx_path}")
        else:
            raise RuntimeError("prefer-basic")
    except Exception as e:
        print(f"[FALLBACK] {e}")
        build_docx_with_python_docx_to_path(
            pages_text,
            images_by_page=ocr_images_by_page,
            crops_by_page=crops_by_page,
            out_path=docx_path,
            insert_page_breaks=not args.no_page_breaks,
            image_max_width_in=args.image_max_width
        )
        print(f"[OK] wrote docx (basic): {docx_path}")

if __name__ == "__main__":
    main()
